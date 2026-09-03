"""Prepare locked Agent 2 inputs from the authoritative source files.

This script performs no substantive statistical modelling. It verifies the raw
workbook checksum, extracts the R1-authorised raw fields to CSV, and preserves
compact source-document extracts used to justify the measurement and scoring
architecture.
"""

# Public repository release.
# This version preserves the analytical logic used in the dissertation.
# Changes from the archived executed version are limited to documentation
# and terminology clarification; no statistical model, scoring formula,
# hypothesis test, exclusion rule, or numerical procedure has been changed.

from __future__ import annotations

import hashlib
import json
import platform
import sys
from pathlib import Path

import pandas as pd
import openpyxl
from docx import Document
import docx


ROOT = Path(__file__).resolve().parents[2]
RAW_XLSX = ROOT / "01_INPUTS" / "03_RAW_DATA" / "BEMM828_Questionnaire Data.xlsx"
QUESTIONNAIRE_DOCX = (
    ROOT
    / "01_INPUTS"
    / "02_QUESTIONNAIRE_AND_CODEBOOK"
    / "BEMM828_Questionnaire_Only_Qualtrics_Build.docx"
)
CHAPTERS_DOCX = ROOT / "01_INPUTS" / "01_CHAPTERS_1_TO_3" / "Chapters 1-3_030926.docx"
OUTDIR = ROOT / "03_ANALYSIS" / "outputs" / "agent_2"
EXPECTED_SHA256 = "3a3bd2e7460ca30053b8f6e741a3dd83621774f988de68da72854021096d299e"

ELIGIBILITY_FIELDS = ["ELIG_AGE", "ELIG_UK", "ELIG_DB12", "ELIG_ENG"]
PRIVACY_ITEMS = [
    "PC_CTRL1",
    "PC_CTRL2",
    "PC_AWA1",
    "PC_AWA2",
    "PC_COLL1",
    "PC_COLL2",
    "PC_COLL3",
    "PC_COLL4",
]
TRANSPARENCY_ITEMS = ["PT_D1", "PT_D2", "PT_C1", "PT_C2", "PT_A1", "PT_A2"]
TRUST_ITEMS = [
    "CT_COMP1",
    "CT_COMP2",
    "CT_COMP3",
    "CT_BEN1",
    "CT_BEN2",
    "CT_BEN3",
    "CT_INT1",
    "CT_INT2",
    "CT_INT3",
]
WTD_ITEMS = ["WTD1", "WTD2", "WTD3", "WTD4"]
DEMOGRAPHICS = ["DEM_AGE", "DEM_GENDER", "DEM_DBFREQ", "DEM_AIUSE"]
AUTHORITATIVE_FIELDS = (
    ["Respondent_ID"] + ELIGIBILITY_FIELDS + PRIVACY_ITEMS + TRANSPARENCY_ITEMS + TRUST_ITEMS + WTD_ITEMS + DEMOGRAPHICS
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def table_payload(path: Path) -> list[dict[str, object]]:
    doc = Document(path)
    tables: list[dict[str, object]] = []
    for index, table in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables.append({"table_index": index, "rows": rows})
    return tables


def chapter_3_extract(path: Path) -> list[str]:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    start = next(i for i, text in enumerate(paragraphs) if text.startswith("Chapter 3"))
    end = next((i for i, text in enumerate(paragraphs[start + 1 :], start + 1) if text == "References"), len(paragraphs))
    return paragraphs[start:end]


def main() -> None:
    OUTDIR.mkdir(parents=True, exist_ok=True)
    observed_sha = sha256(RAW_XLSX)
    if observed_sha != EXPECTED_SHA256:
        raise SystemExit(f"Raw workbook checksum mismatch: {observed_sha}")

    df = pd.read_excel(RAW_XLSX, sheet_name="Data")
    missing = [col for col in AUTHORITATIVE_FIELDS if col not in df.columns]
    if missing:
        raise SystemExit(f"Authoritative fields missing from raw workbook: {missing}")
    locked = df[AUTHORITATIVE_FIELDS].copy()
    locked.to_csv(OUTDIR / "agent_2_locked_raw_fields.csv", index=False, encoding="utf-8")

    source_extract = {
        "raw_workbook": str(RAW_XLSX.relative_to(ROOT)),
        "expected_sha256": EXPECTED_SHA256,
        "observed_sha256": observed_sha,
        "checksum_match": True,
        "questionnaire_tables": table_payload(QUESTIONNAIRE_DOCX),
        "chapter_3_paragraphs": chapter_3_extract(CHAPTERS_DOCX),
        "authoritative_fields_written": AUTHORITATIVE_FIELDS,
        "excluded_workbook_material": [
            "PRIVACY_TOTAL",
            "PRIVACY_MEAN",
            "TRANSPARENCY_TOTAL",
            "TRANSPARENCY_MEAN",
            "TRUST_TOTAL",
            "TRUST_MEAN",
            "WTD_TOTAL",
            "WTD_MEAN",
            "Scale_Checks sheet",
        ],
    }
    (OUTDIR / "agent_2_source_extract.json").write_text(
        json.dumps(source_extract, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    prepare_session = {
        "python_executable": sys.executable,
        "python_version": sys.version,
        "platform": platform.platform(),
        "pandas_version": pd.__version__,
        "openpyxl_version": openpyxl.__version__,
        "python_docx_version": docx.__version__,
        "purpose": "checksum verification, raw XLSX extraction, and DOCX source extraction only",
    }
    (OUTDIR / "input_extraction_session_info.json").write_text(
        json.dumps(prepare_session, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    print(json.dumps({"status": "PASS", "output_dir": str(OUTDIR), "sha256": observed_sha}, indent=2))


if __name__ == "__main__":
    main()
