"""Agent 1 reproducible data audit for BEMM828 Chapter 5.

This script verifies workbook integrity, schema, codebook coverage, response
coding, missingness, and permitted diagnostic data-quality flags. It does not
estimate hypothesis relationships, mediation, structural paths, reliability, or
construct composites for substantive analysis.
"""

# Public repository release.
# This version preserves the analytical logic used in the dissertation.
# Changes from the archived executed version are limited to documentation
# and terminology clarification; no statistical model, scoring formula,
# hypothesis test, exclusion rule, or numerical procedure has been changed.

from __future__ import annotations

import hashlib
import json
import re
from collections import Counter
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[2]
RAW_XLSX = ROOT / "01_INPUTS" / "03_RAW_DATA" / "BEMM828_Questionnaire Data.xlsx"
QUESTIONNAIRE_DOCX = (
    ROOT
    / "01_INPUTS"
    / "02_QUESTIONNAIRE_AND_CODEBOOK"
    / "BEMM828_Questionnaire_Only_Qualtrics_Build.docx"
)
EXPECTED_SHA256 = "3a3bd2e7460ca30053b8f6e741a3dd83621774f988de68da72854021096d299e"

ELIGIBILITY_FIELDS = ["ELIG_AGE", "ELIG_UK", "ELIG_DB12", "ELIG_ENG"]
PRIVACY_ITEMS = [
    "PC_CTRL1",
    "PC_CTRL2",
    "PC_AWA1",
    "PC_AWA2",
    "PC_COLL1",
    "PC_COLL2",
    "PC_COLL3",
    "PC_COLL4",
]
TRANSPARENCY_ITEMS = ["PT_D1", "PT_D2", "PT_C1", "PT_C2", "PT_A1", "PT_A2"]
TRUST_ITEMS = [
    "CT_COMP1",
    "CT_COMP2",
    "CT_COMP3",
    "CT_BEN1",
    "CT_BEN2",
    "CT_BEN3",
    "CT_INT1",
    "CT_INT2",
    "CT_INT3",
]
WTD_ITEMS = ["WTD1", "WTD2", "WTD3", "WTD4"]
FOCAL_ITEMS = PRIVACY_ITEMS + TRANSPARENCY_ITEMS + TRUST_ITEMS + WTD_ITEMS
DEMOGRAPHICS = ["DEM_AGE", "DEM_GENDER", "DEM_DBFREQ", "DEM_AIUSE"]
AUTHORITATIVE_RAW_FIELDS = ["Respondent_ID"] + ELIGIBILITY_FIELDS + FOCAL_ITEMS + DEMOGRAPHICS
DERIVED_NON_AUTHORITATIVE = [
    "PRIVACY_TOTAL",
    "PRIVACY_MEAN",
    "TRANSPARENCY_TOTAL",
    "TRANSPARENCY_MEAN",
    "TRUST_TOTAL",
    "TRUST_MEAN",
    "WTD_TOTAL",
    "WTD_MEAN",
]

EXPECTED_ALLOWED_VALUES = {
    "privacy_7_point": set(range(1, 8)),
    "five_point": set(range(1, 6)),
    "eligibility": {"Yes", "No"},
    "DEM_AGE": {
        "18-24",
        "18–24",
        "25-34",
        "25–34",
        "35-44",
        "35–44",
        "45-54",
        "45–54",
        "55-64",
        "55–64",
        "65 or over",
        "Prefer not to say",
    },
    "DEM_GENDER": {
        "Man",
        "Woman",
        "Non-binary or another gender identity",
        "Prefer not to say",
    },
    "DEM_DBFREQ": {
        "Every day",
        "Several times a week",
        "About once a week",
        "A few times a month",
        "Less often",
        "Prefer not to say",
    },
    "DEM_AIUSE": {"Yes", "No", "Not sure", "Prefer not to say"},
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def workbook_schema(path: Path) -> dict[str, Any]:
    wb = load_workbook(path, read_only=True, data_only=False)
    wb_values = load_workbook(path, read_only=True, data_only=True)
    schema: dict[str, Any] = {}
    for ws in wb.worksheets:
        headers = [cell.value for cell in next(ws.iter_rows(min_row=1, max_row=1))]
        value_ws = wb_values[ws.title]
        displayed_error_cells = 0
        formula_cells = 0
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.startswith("="):
                    formula_cells += 1
        for row in value_ws.iter_rows(values_only=True):
            displayed_error_cells += sum(1 for value in row if isinstance(value, str) and value.startswith("#"))
        schema[ws.title] = {
            "rows_including_header": ws.max_row,
            "data_rows": max(ws.max_row - 1, 0),
            "columns": ws.max_column,
            "headers": headers,
            "formula_cells": formula_cells,
            "displayed_error_cells": displayed_error_cells,
        }
    return schema


def extract_codebook_ids(path: Path) -> dict[str, Any]:
    doc = Document(path)
    table_ids: list[str] = []
    table_shapes: list[dict[str, int]] = []
    for table in doc.tables:
        table_shapes.append({"rows": len(table.rows), "columns": len(table.columns)})
        for row in table.rows[1:]:
            if row.cells:
                value = row.cells[0].text.strip()
                if value:
                    table_ids.append(value)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return {
        "table_count": len(doc.tables),
        "table_shapes": table_shapes,
        "variable_ids": table_ids,
        "mentions_explicit_consent": any("consent" in p.lower() for p in paragraphs),
        "mentions_completion_time": any("completion time" in p.lower() for p in paragraphs),
    }


def value_counts(series: pd.Series) -> dict[str, int]:
    counts = series.fillna("<MISSING>").astype(str).value_counts(dropna=False)
    return {str(k): int(v) for k, v in counts.items()}


def coding_audit(df: pd.DataFrame) -> dict[str, Any]:
    checks: dict[str, Any] = {}
    for name, fields, allowed in [
        ("privacy_7_point", PRIVACY_ITEMS, EXPECTED_ALLOWED_VALUES["privacy_7_point"]),
        ("transparency_5_point", TRANSPARENCY_ITEMS, EXPECTED_ALLOWED_VALUES["five_point"]),
        ("trust_5_point", TRUST_ITEMS, EXPECTED_ALLOWED_VALUES["five_point"]),
        ("wtd_5_point", WTD_ITEMS, EXPECTED_ALLOWED_VALUES["five_point"]),
    ]:
        present_values = sorted({int(v) for v in df[fields].stack().dropna().unique()})
        invalid = int((~df[fields].stack().dropna().isin(allowed)).sum())
        checks[name] = {
            "fields": fields,
            "allowed_values": sorted(allowed),
            "observed_values": present_values,
            "invalid_cells": invalid,
            "missing_cells": int(df[fields].isna().sum().sum()),
            "dtypes": {col: str(df[col].dtype) for col in fields},
        }
    for col in ELIGIBILITY_FIELDS:
        observed = set(df[col].dropna().astype(str))
        checks[col] = {
            "allowed_values": sorted(EXPECTED_ALLOWED_VALUES["eligibility"]),
            "observed_values": sorted(observed),
            "invalid_cells": int((~df[col].dropna().astype(str).isin(EXPECTED_ALLOWED_VALUES["eligibility"])).sum()),
            "missing_cells": int(df[col].isna().sum()),
        }
    for col in DEMOGRAPHICS:
        observed = set(df[col].dropna().astype(str))
        checks[col] = {
            "allowed_values": sorted(EXPECTED_ALLOWED_VALUES[col]),
            "observed_values": sorted(observed),
            "invalid_cells": int((~df[col].dropna().astype(str).isin(EXPECTED_ALLOWED_VALUES[col])).sum()),
            "missing_cells": int(df[col].isna().sum()),
        }
    return checks


def diagnostic_flags(df: pd.DataFrame) -> dict[str, Any]:
    duplicate_id_rows = int(df["Respondent_ID"].duplicated(keep=False).sum())
    duplicate_focal_vectors = df.duplicated(subset=FOCAL_ITEMS, keep=False)
    duplicate_groups = Counter(tuple(row) for row in df.loc[duplicate_focal_vectors, FOCAL_ITEMS].to_numpy().tolist())

    straightline = {
        "privacy_block_all_8_same": df[PRIVACY_ITEMS].nunique(axis=1, dropna=False).eq(1),
        "transparency_block_all_6_same": df[TRANSPARENCY_ITEMS].nunique(axis=1, dropna=False).eq(1),
        "trust_block_all_9_same": df[TRUST_ITEMS].nunique(axis=1, dropna=False).eq(1),
        "wtd_block_all_4_same": df[WTD_ITEMS].nunique(axis=1, dropna=False).eq(1),
        "all_27_focal_items_same_numeric_response": df[FOCAL_ITEMS].nunique(axis=1, dropna=False).eq(1),
    }
    any_construct_straightline = (
        straightline["privacy_block_all_8_same"]
        | straightline["transparency_block_all_6_same"]
        | straightline["trust_block_all_9_same"]
        | straightline["wtd_block_all_4_same"]
    )
    return {
        "rules_are_diagnostic_only": True,
        "duplicate_respondent_id_rows": duplicate_id_rows,
        "duplicate_focal_response_vector_rows": int(duplicate_focal_vectors.sum()),
        "duplicate_focal_response_vector_groups": len(duplicate_groups),
        "straightline_counts": {name: int(mask.sum()) for name, mask in straightline.items()},
        "any_construct_block_straightline_rows": int(any_construct_straightline.sum()),
    }


def main() -> None:
    file_hash = sha256(RAW_XLSX)
    schema = workbook_schema(RAW_XLSX)
    codebook = extract_codebook_ids(QUESTIONNAIRE_DOCX)
    df = pd.read_excel(RAW_XLSX, sheet_name="Data")

    expected_codebook_ids = ELIGIBILITY_FIELDS + FOCAL_ITEMS + DEMOGRAPHICS
    workbook_columns = list(df.columns)
    missing_expected_authoritative = [c for c in AUTHORITATIVE_RAW_FIELDS if c not in workbook_columns]
    missing_derived = [c for c in DERIVED_NON_AUTHORITATIVE if c not in workbook_columns]
    unexpected_data_columns = [c for c in workbook_columns if c not in AUTHORITATIVE_RAW_FIELDS + DERIVED_NON_AUTHORITATIVE]
    timing_pattern = re.compile(r"(^|_)(duration|completion|start|end|timestamp|recorded|finished|progress|time)($|_)")
    timing_like_columns = [
        c
        for c in workbook_columns
        if timing_pattern.search(c.lower())
    ]
    consent_like_columns = [c for c in workbook_columns if "consent" in c.lower()]

    focal_complete_mask = df[FOCAL_ITEMS].notna().all(axis=1)
    all_eligible_mask = df[ELIGIBILITY_FIELDS].eq("Yes").all(axis=1)
    summary = {
        "raw_file": str(RAW_XLSX.relative_to(ROOT)),
        "sha256": file_hash,
        "sha256_matches_expected": file_hash == EXPECTED_SHA256,
        "workbook_schema": schema,
        "data_dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
        "codebook_table_variable_ids": codebook,
        "mapping_audit": {
            "expected_codebook_ids": expected_codebook_ids,
            "codebook_ids_missing_from_workbook": [c for c in expected_codebook_ids if c not in workbook_columns],
            "workbook_ids_missing_from_codebook": [c for c in expected_codebook_ids if c not in codebook["variable_ids"]],
            "missing_expected_authoritative_fields": missing_expected_authoritative,
            "derived_non_authoritative_columns_present": [c for c in DERIVED_NON_AUTHORITATIVE if c in workbook_columns],
            "derived_non_authoritative_columns_missing": missing_derived,
            "unexpected_data_columns": unexpected_data_columns,
            "consent_like_columns": consent_like_columns,
            "timing_like_columns": timing_like_columns,
        },
        "sample_flow_boundary": {
            "records_in_data_sheet": int(len(df)),
            "all_four_eligibility_yes_rows": int(all_eligible_mask.sum()),
            "rows_with_any_non_yes_eligibility": int((~all_eligible_mask).sum()),
            "consent_in_questionnaire_build_note": bool(codebook["mentions_explicit_consent"]),
            "consent_column_in_workbook": bool(consent_like_columns),
            "survey_starts_or_screenouts_directly_supported": False,
        },
        "eligibility_value_counts": {col: value_counts(df[col]) for col in ELIGIBILITY_FIELDS},
        "demographic_value_counts": {col: value_counts(df[col]) for col in DEMOGRAPHICS},
        "missingness": {
            "focal_item_missing_cells": int(df[FOCAL_ITEMS].isna().sum().sum()),
            "focal_item_missing_by_column": {col: int(df[col].isna().sum()) for col in FOCAL_ITEMS},
            "complete_focal_rows": int(focal_complete_mask.sum()),
            "incomplete_focal_rows": int((~focal_complete_mask).sum()),
            "authoritative_raw_missing_cells_by_column": {
                col: int(df[col].isna().sum()) for col in AUTHORITATIVE_RAW_FIELDS if col in df.columns
            },
        },
        "coding_audit": coding_audit(df),
        "diagnostic_quality_flags": diagnostic_flags(df),
        "non_authoritative_workbook_material": {
            "scale_checks_sheet_present": "Scale_Checks" in schema,
            "data_sheet_derived_columns_present": [c for c in DERIVED_NON_AUTHORITATIVE if c in workbook_columns],
        },
    }
    print(json.dumps(summary, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
